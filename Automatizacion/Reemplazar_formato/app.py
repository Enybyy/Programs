import os
import io
import pandas as pd
from flask import Flask, request, render_template, send_file, redirect, flash
from docx import Document
from werkzeug.utils import secure_filename
import zipfile
from datetime import datetime
from docx2pdf import convert
import tempfile
import shutil

app = Flask(__name__)
app.secret_key = 'secret_key_for_session'  # Cambia esto por una clave segura en producción

# Carpetas de almacenamiento temporal
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'generated'
ALLOWED_EXTENSIONS = {'xlsx', 'docx'}

# Crear las carpetas si no existen
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename, ext):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == ext

def replace_text_in_paragraph(paragraph, replacements):
    """Reemplaza etiquetas en el texto de un párrafo."""
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        # Validar que se hayan enviado los tres archivos: DB-Contratos, DB-Personal y plantilla Word
        if ('contracts_file' not in request.files or 
            'personal_file' not in request.files or 
            'template_file' not in request.files):
            flash('Falta algún archivo.', 'danger')
            return redirect(request.url)
        
        contracts_file = request.files['contracts_file']
        personal_file = request.files['personal_file']
        template_file = request.files['template_file']
        
        if (contracts_file.filename == '' or 
            personal_file.filename == '' or 
            template_file.filename == ''):
            flash('No se seleccionó alguno de los archivos.', 'danger')
            return redirect(request.url)
        
        if not (allowed_file(contracts_file.filename, 'xlsx') and 
                allowed_file(personal_file.filename, 'xlsx') and 
                allowed_file(template_file.filename, 'docx')):
            flash('Tipo de archivo no válido. Asegúrate de subir dos archivos Excel (.xlsx) y una plantilla Word (.docx).', 'danger')
            return redirect(request.url)
        
        # Guardar los archivos de forma segura
        contracts_filename = secure_filename(contracts_file.filename)
        personal_filename = secure_filename(personal_file.filename)
        template_filename = secure_filename(template_file.filename)
        
        contracts_path = os.path.join(UPLOAD_FOLDER, contracts_filename)
        personal_path = os.path.join(UPLOAD_FOLDER, personal_filename)
        template_path = os.path.join(UPLOAD_FOLDER, template_filename)
        
        contracts_file.save(contracts_path)
        personal_file.save(personal_path)
        template_file.save(template_path)
        
        # Crear un directorio de salida único con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        output_dir = os.path.join(OUTPUT_FOLDER, f"output_{timestamp}")
        os.makedirs(output_dir, exist_ok=True)
        
        # Crear subcarpetas para Contratos y Bases de Datos dentro del output_dir
        contratos_folder = os.path.join(output_dir, "Contratos")
        bases_folder = os.path.join(output_dir, "Bases de Datos")
        os.makedirs(contratos_folder, exist_ok=True)
        os.makedirs(bases_folder, exist_ok=True)
        
        # Cargar las dos bases de datos Excel
        try:
            df_contracts = pd.read_excel(contracts_path, dtype=str)
            df_personal = pd.read_excel(personal_path, dtype=str)
            total_records = len(df_contracts)
            flash(f'Se han detectado {total_records} registros en la base de contratos.', 'info')
        except Exception as e:
            flash(f'Error al leer alguno de los archivos Excel: {e}', 'danger')
            return redirect(request.url)
        
        # Convertir las columnas de fecha en la DB-Contratos (si existen)
        if 'Fecha de inicio' in df_contracts.columns:
            df_contracts['Fecha de inicio'] = pd.to_datetime(df_contracts['Fecha de inicio'], errors='coerce', dayfirst=True)
        if 'Fecha de fin' in df_contracts.columns:
            df_contracts['Fecha de fin'] = pd.to_datetime(df_contracts['Fecha de fin'], errors='coerce', dayfirst=True)
        
        # Realizar el merge de las dos bases de datos usando "Número de documento" como clave
        # Se asume que ambas bases tienen la columna "Número de documento"
        df_merged = pd.merge(df_contracts, df_personal, on="Número de documento", how="left", suffixes=("", " verificado"))
        
        # Reordenar las columnas según el orden requerido:
        # 1. Tipo de documento (DB-Contratos)
        # 2. Número de documento
        # 3. Apellidos y nombres (original de DB-Contratos)
        # 4. Apellidos y nombres verificados (de DB-Personal)
        # 5. Dirección (DB-Personal)
        # 6. Distrito (DB-Personal)
        # 7. Ciudad (DB-Personal)
        # 8. Campaña (DB-Contratos)
        # 9. Cargo (DB-Contratos)
        # 10. Fecha de inicio (DB-Contratos)
        # 11. Fecha de fin (DB-Contratos)
        # 12. Pago (DB-Contratos)
        desired_order = [
            "Tipo de documento", "Número de documento", "Apellidos y nombres",
            "Apellidos y nombres verificados", "Dirección", "Distrito", "Ciudad",
            "Campaña", "Cargo", "Fecha de inicio", "Fecha de fin", "Pago"
        ]
        # Es posible que en df_merged algunas columnas no existan exactamente con esos nombres.
        # Asegurarse de que "Apellidos y nombres verificados", "Dirección", "Distrito" y "Ciudad" provengan de la DB-Personal.
        # Si las columnas en df_personal tienen nombres distintos, ajusta aquí.
        df_merged = df_merged.reindex(columns=desired_order)
        
        # Separar registros que tuvieron coincidencia y los que no
        df_coinciden = df_merged[df_merged["Apellidos y nombres verificados"].notna()].copy()
        df_no_coinciden = df_merged[df_merged["Apellidos y nombres verificados"].isna()].copy()
        
        # Guardar los resultados del merge en Excel dentro de la carpeta "Bases de Datos"
        merge_file_path = os.path.join(bases_folder, "DB_Coinciden.xlsx")
        no_match_file_path = os.path.join(bases_folder, "DB_No_Coinciden.xlsx")
        try:
            df_coinciden.to_excel(merge_file_path, index=False)
            df_no_coinciden.to_excel(no_match_file_path, index=False)
        except Exception as e:
            flash(f"Error al guardar las bases de datos: {e}", "danger")
            return redirect(request.url)
        
        # Generar los contratos (solo para registros que tienen coincidencia)
        processed_files = 0
        for _, row in df_coinciden.iterrows():
            # Crear el documento Word a partir de la plantilla
            doc = Document(template_path)
            
            # Formatear fechas si son datetime (si no, dejar en blanco)
            fecha_inicio = row["Fecha de inicio"].strftime('%d/%m/%Y') if pd.notna(row.get("Fecha de inicio")) and not isinstance(row["Fecha de inicio"], str) else row.get("Fecha de inicio", "")
            fecha_fin = row["Fecha de fin"].strftime('%d/%m/%Y') if pd.notna(row.get("Fecha de fin")) and not isinstance(row["Fecha de fin"], str) else row.get("Fecha de fin", "")
            
            # Diccionario de reemplazo para la plantilla
            replacements = {
                "[TIPO_DOCUMENTO]": row.get("Tipo de documento", ""),
                "[NUMERO_DOCUMENTO]": row.get("Número de documento", ""),
                "[NOMBRE]": row.get("Apellidos y nombres verificados", ""),  # Usamos el nombre verificado para el contrato
                "[DIRECCION]": row.get("Dirección", ""),
                "[DISTRITO]": row.get("Distrito", ""),
                "[CIUDAD]": row.get("Ciudad", ""),
                "[CAMPANA]": row.get("Campaña", ""),
                "[CARGO]": row.get("Cargo", ""),
                "[FECHA_DE_INICIO]": fecha_inicio,
                "[FECHA_DE_FIN]": fecha_fin,
                "[PAGO]": row.get("Pago", "")
            }
            
            # Reemplazar en párrafos
            for paragraph in doc.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            # Reemplazar en tablas, si las hubiera
            for table in doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        for paragraph in cell.paragraphs:
                            replace_text_in_paragraph(paragraph, replacements)
            
            # Nombre del archivo: usar el nombre verificado con espacios (sin modificar)
            full_name = row.get("Apellidos y nombres verificados", "nombre no definido")
            word_filename = f"{full_name}.docx"
            word_file_path = os.path.join(contratos_folder, word_filename)
            doc.save(word_file_path)
            
            # Convertir el archivo Word a PDF usando docx2pdf
            # Utilizaremos un directorio temporal para la conversión
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_word = os.path.join(temp_dir, "temp.docx")
                temp_pdf = os.path.join(temp_dir, "temp.pdf")
                shutil.copy(word_file_path, temp_word)
                try:
                    convert(temp_word, temp_pdf)
                    # Mover el PDF al folder de Contratos con el mismo nombre que el Word
                    pdf_filename = f"{full_name}.pdf"
                    pdf_file_path = os.path.join(contratos_folder, pdf_filename)
                    shutil.move(temp_pdf, pdf_file_path)
                except Exception as e:
                    flash(f"Error al convertir a PDF para {full_name}: {e}", "warning")
            
            processed_files += 1
        
        flash(f'Se han generado {processed_files} contratos con éxito.', 'success')
        
        # Empaquetar toda la carpeta output_dir en un archivo ZIP
        zip_stream = io.BytesIO()
        with zipfile.ZipFile(zip_stream, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for foldername, subfolders, filenames in os.walk(output_dir):
                for filename in filenames:
                    filepath = os.path.join(foldername, filename)
                    arcname = os.path.relpath(filepath, output_dir)
                    zipf.write(filepath, arcname)
        zip_stream.seek(0)
        return send_file(zip_stream, as_attachment=True, download_name=f"contratos_{timestamp}.zip", mimetype='application/zip')
    
    # Renderizamos la plantilla index.html (frontend no se modifica)
    return render_template('index.html')

if __name__ == '__main__':
    app.run(debug=True)
