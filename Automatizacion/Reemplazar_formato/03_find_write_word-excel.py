import os
import pandas as pd
from docx import Document

# Configuración de rutas
excel_file = 'data/db.xlsx'           # Ruta del archivo Excel
template_file = 'data/formato.docx'         # Ruta de la plantilla de Word
output_folder = 'data/creado/'                # Carpeta donde se guardarán los contratos

# Crear la carpeta de salida si no existe
if not os.path.exists(output_folder):
    os.makedirs(output_folder)

# Leer la base de datos de Excel y convertir fechas
df = pd.read_excel(excel_file, dtype=str)  # Leer todo como str para evitar errores
df['FECHA DE INICIO'] = pd.to_datetime(df['FECHA DE INICIO'], errors='coerce', dayfirst=True)
df['FECHA DE FIN'] = pd.to_datetime(df['FECHA DE FIN'], errors='coerce', dayfirst=True)

def replace_text_in_paragraph(paragraph, replacements):
    """ Reemplaza texto en un párrafo """
    for key, value in replacements.items():
        if key in paragraph.text:
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(value))

# Iterar sobre cada registro del Excel
for index, row in df.iterrows():
    doc = Document(template_file)
    
    # Convertir fechas a formato dd/mm/yyyy si no son NaT (Not a Time)
    fecha_inicio = row['FECHA DE INICIO'].strftime('%d/%m/%Y') if pd.notna(row['FECHA DE INICIO']) else ''
    fecha_fin = row['FECHA DE FIN'].strftime('%d/%m/%Y') if pd.notna(row['FECHA DE FIN']) else ''
    
    # Diccionario de reemplazo
    replacements = {
        '[NOMBRE]': row['NOMBRE'],
        '[APELLIDO]': row['APELLIDO'],
        '[DNI]': row['DNI'],
        '[UBICACION]': row['UBICACION'],
        '[CAMPAÑA]': row['CAMPAÑA'],
        '[FECHA DE INICIO]': fecha_inicio,
        '[FECHA DE FIN]': fecha_fin,
        '[SALARIO]': row['SALARIO']
    }
    
    # Reemplazar etiquetas en el documento
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # También recorrer tablas si hay celdas con etiquetas
    for table in doc.tables:
        for row_table in table.rows:
            for cell in row_table.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)

    # Guardar el contrato con nombre personalizado
    nombre_archivo = f"{row['NOMBRE']}_{row['APELLIDO'].replace(' ', '_')}.docx"
    output_path = os.path.join(output_folder, nombre_archivo)
    doc.save(output_path)
    print(f"Contrato generado: {output_path}")